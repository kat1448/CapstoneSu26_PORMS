from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Report3_SRS_Updated_20260731.docx"
DOWNLOAD = Path(r"C:\Users\LENOVO\Downloads\Report3_SRS_Updated_20260731.docx")

BLUE = "164A72"
LIGHT_BLUE = "EAF3F9"
MID_BLUE = "2E78AD"
GREY = "5F7180"
LIGHT_GREY = "F4F7F9"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, WHITE, 9)
        shade(table.rows[0].cells[i], BLUE)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, None, font_size)
            if ri % 2:
                shade(cells[i], LIGHT_GREY)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_requirement(doc, code, title, trigger, description, normal, exceptions, rules, route):
    add_heading(doc, f"{code} - {title}", 3)
    add_table(doc, ["Item", "Specification"], [
        ("Screen / route", route),
        ("Trigger", trigger),
        ("Description", description),
        ("Normal flow", normal),
        ("Alternative / exception flow", exceptions),
        ("Related rules", rules),
    ], [3.5, 13.5], 9)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and select Update Field to refresh the table of contents."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, end])


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("PORMS SRS  |  ")
    run.font.size = Pt(8)
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.8); sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.0)
sec.header_distance = Cm(0.8); sec.footer_distance = Cm(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size, color in [("Title", 26, BLUE), ("Heading 1", 18, BLUE), ("Heading 2", 15, MID_BLUE), ("Heading 3", 12, BLUE), ("Heading 4", 11, GREY)]:
    styles[name].font.name = "Arial"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor.from_string(color)
    styles[name].font.bold = True
styles["Heading 1"].paragraph_format.space_before = Pt(16)
styles["Heading 2"].paragraph_format.space_before = Pt(12)

header = sec.header.paragraphs[0]
header.text = "PORMS — PORT OPERATION RISK MANAGEMENT SYSTEM"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Arial"; header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor.from_string(GREY)
page_number(sec.footer.paragraphs[0])

# Cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(60)
r = p.add_run("FPT UNIVERSITY DA NANG"); r.bold = True; r.font.name = "Arial"; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(55)
r = p.add_run("PORMS"); r.bold = True; r.font.name = "Arial"; r.font.size = Pt(38); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PORT OPERATION RISK MANAGEMENT SYSTEM"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(MID_BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(30)
r = p.add_run("REPORT 3\nSOFTWARE REQUIREMENTS SPECIFICATION"); r.bold = True; r.font.size = Pt(22)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(45)
p.add_run("Document version: 2.0\nCode baseline: 0001_data_fullflow / local working baseline\nUpdated: 31 July 2026").font.size = Pt(11)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(80)
p.add_run("Da Nang, July 2026").italic = True
doc.add_page_break()

add_heading(doc, "Document Control", 1)
add_table(doc, ["Field", "Value"], [
    ("Document", "Report 3 — Software Requirements Specification"),
    ("Project", "PORMS — Port Operation Risk Management System"),
    ("Version", "2.0"),
    ("Status", "Updated to match the implemented July 2026 baseline"),
    ("Primary scope", "Weather ingestion, risk evaluation, operational alerts and tasks, forecasting, simulation, reporting and audit"),
    ("Deployment", "Full Docker stack: React frontend, ASP.NET Core API, PostgreSQL, Prefect and Metabase"),
], [4, 13])

add_heading(doc, "Record of Changes", 1)
add_table(doc, ["Date", "Type", "Owner", "Change description"], [
    ("04–06/05/2026", "A", "Project team", "Initial SRS covering authentication, port, dashboard, alert, risk, SOP, simulation and analytics."),
    ("29/05–05/06/2026", "M/A", "Project team", "Updated role naming, diagrams, screen flows, ERD and document structure."),
    ("31/07/2026", "M", "PORMS team", "Reconciled requirements with the current source code, database schema, routes and Docker deployment."),
    ("31/07/2026", "A", "PORMS team", "Added port-scoped alert delivery, task assignment lifecycle, SMTP email, Vietnamese voice alert, notification preferences, operational forecast, AI long-range trend, forecast evaluation and Excel/PDF reporting."),
    ("31/07/2026", "R", "PORMS team", "Removed or reclassified unsupported legacy claims such as implemented forgot-password flow, automatic lockout after five attempts, obsolete task_logs/alert_recipients entities and old role permissions."),
], [2.5, 1.5, 3.2, 9.8])
doc.add_page_break()

add_heading(doc, "Table of Contents", 1)
add_toc(doc.add_paragraph())
doc.add_page_break()

add_heading(doc, "1. Introduction", 1)
add_heading(doc, "1.1 Purpose", 2)
add_body(doc, "This Software Requirements Specification defines the functional and non-functional requirements of PORMS at the July 2026 implementation baseline. It is intended for the project team, supervisors, testers and evaluators. The document separates user-visible requirements from automated backend and ETL behavior and aligns terminology with the Vietnamese web interface.")
add_heading(doc, "1.2 Product Scope", 2)
add_body(doc, "PORMS is a decision-support system for port operation safety. It collects weather information, evaluates operational risk, recommends an operating mode, creates alerts and response tasks, supports controlled simulation, provides short- and long-range planning views, evaluates forecast accuracy, and exports auditable operational reports.")
add_bullets(doc, [
    "Operational scope: multiple ports and operational zones, demonstrated with Tiên Sa and additional Vietnamese ports.",
    "Risk scope: wind, rainfall and visibility aggregated into LOW, MEDIUM, HIGH and CRITICAL.",
    "Decision scope: NORMAL, LIMITED and STOP operation modes with SOP-based response recommendations.",
    "Accountability scope: role- and port-based access, assignment to a named Operator, completion notes and operation-event history.",
    "Planning scope: OpenWeather five-day forecast, data-driven long-range trend analysis and forecast-versus-actual evaluation.",
])
add_heading(doc, "1.3 Definitions and Abbreviations", 2)
add_table(doc, ["Term", "Definition"], [
    ("PORMS", "Port Operation Risk Management System."),
    ("SOP", "Standard Operating Procedure used to recommend or trigger a response to risk."),
    ("ETL", "Extract, Transform and Load pipeline orchestrated by Prefect."),
    ("RBAC", "Role-Based Access Control."),
    ("Port Manager", "Manager assigned to one port and responsible for responding to that port's alerts."),
    ("Operator", "Operational employee assigned to one port and able to process tasks assigned personally."),
    ("Actual data", "Observed weather reading available after a forecast time has elapsed."),
    ("Simulation", "Controlled replay of stored or prepared weather snapshots through the risk workflow."),
], [3.4, 13.6])
add_heading(doc, "1.4 Sources of Truth", 2)
add_bullets(doc, [
    "Frontend routes and permissions: frontend/src/router/index.tsx and frontend/src/navigation/navigation.ts.",
    "REST behavior: backend/PORMS.API/Controllers.",
    "Access-scoped data queries: backend/PORMS.Infrastructure/Repositories.",
    "Database definition: docs/database/schema.sql.",
    "Runtime topology: infra/docker-compose.yml and Dockerfiles.",
    "Automated ETL: etl flows and Prefect deployments.",
])

add_heading(doc, "2. Product Overview", 1)
add_heading(doc, "2.1 System Context", 2)
add_table(doc, ["Source / actor", "Interaction", "PORMS output"], [
    ("OpenWeather API", "Provides current weather and five-day forecast data.", "Normalized weather readings and forecast cards."),
    ("Prefect worker", "Runs scheduled weather collection, transformation, historical backfill and analytics loading.", "Operational and analytics records."),
    ("Admin", "Manages users, ports, zones, simulation and system-wide monitoring.", "Configuration and audit events."),
    ("Port Manager", "Monitors one assigned port, receives High/Critical alerts and assigns response tasks.", "Acknowledgement, assignment and oversight records."),
    ("Operator", "Receives and executes personally assigned tasks.", "Acknowledgement, progress and completion notes."),
    ("SMTP server", "Delivers alert and assignment emails.", "Delivery attempt recorded in application log/response metadata."),
    ("Metabase", "Reads the analytics schema.", "Historical KPI and BI dashboards."),
], [3.4, 6.7, 6.9])

add_heading(doc, "2.2 Technology Baseline", 2)
add_table(doc, ["Layer", "Technology", "Responsibility"], [
    ("Web UI", "React 18, TypeScript 5, Vite 5, MUI/X Charts, Leaflet", "Responsive Vietnamese operational interface."),
    ("API", "ASP.NET Core 10", "Authentication, authorization, workflows, reports and integrations."),
    ("Database", "PostgreSQL 16", "Operational and analytics schemas."),
    ("ETL", "Python 3.11, Prefect 2", "Scheduled weather ingestion and analytics loading."),
    ("BI", "Metabase 0.50", "Historical dashboards."),
    ("Deployment", "Docker Compose", "Runs PostgreSQL, backend, frontend, Prefect server/worker and Metabase."),
], [3, 5, 9])

add_heading(doc, "2.3 End-to-End Operational Flow", 2)
add_table(doc, ["Step", "Processing", "Result"], [
    ("1", "Collect current weather for each active port.", "A normalized weather reading is stored."),
    ("2", "Evaluate wind, rain and visibility thresholds.", "A final risk level is produced."),
    ("3", "Apply matching SOP and recommended operating mode.", "NORMAL, LIMITED or STOP guidance is produced."),
    ("4", "If risk reaches HIGH or CRITICAL, create an alert and response task.", "Port-scoped operational records are stored."),
    ("5", "Notify the responsible Port Manager.", "Popup/voice and SMTP email are available according to severity and settings."),
    ("6", "Port Manager assigns the task to an eligible Operator at the same port.", "The selected Operator receives task email and gains access to the related alert."),
    ("7", "Operator acknowledges, starts and completes the task with a result note.", "Progress and completion are persisted in the audit history."),
    ("8", "Managers preview and export filtered data.", "Excel/PDF report and report-export event are created."),
], [1.3, 9.5, 6.2])

add_heading(doc, "3. User Requirements", 1)
add_heading(doc, "3.1 Actors and Responsibilities", 2)
add_table(doc, ["Actor", "Scope", "Responsibilities"], [
    ("ADMIN", "All ports; no assigned_port_id required", "Manage users, ports and zones; monitor all alerts/tasks; configure rules; run simulations; use forecasts, evaluation and reports."),
    ("PORT_MANAGER", "Exactly one assigned port", "Monitor that port, receive High/Critical alerts, acknowledge alerts, assign tasks only to Operators at the same port, supervise work, use planning and reports for the same port."),
    ("OPERATOR", "Exactly one assigned port; personal task scope", "View operational overview; view only alerts related to personally assigned tasks; acknowledge/start/complete personal tasks and enter completion notes."),
    ("SYSTEM/ETL", "Configured active ports", "Collect weather, evaluate risk, generate records, refresh analytics and retain audit data."),
], [3, 4, 10])

add_heading(doc, "3.2 Authorization Matrix", 2)
add_table(doc, ["Function", "Admin", "Port Manager", "Operator"], [
    ("Dashboard", "All ports", "Own port", "Own operational scope"),
    ("Alert center/detail", "All alerts", "Own port", "Only alerts linked to assigned tasks"),
    ("Assign task", "Any eligible port Operator", "Operator at own port", "No"),
    ("Acknowledge/start/complete task", "Monitor", "Monitor", "Own assigned task"),
    ("User management", "Create/edit/delete non-Admin accounts", "No", "No"),
    ("Port and zone management", "Yes", "No in current UI", "No"),
    ("Risk threshold and SOP", "Yes", "Yes", "No"),
    ("Simulation", "Yes", "No in current UI", "No"),
    ("Forecast and evaluation", "Yes", "Own port", "No"),
    ("Operational reports", "All ports", "Own port", "No"),
    ("Operation history", "All accessible events", "Own port", "Own accessible events"),
], [6.5, 3.5, 3.5, 3.5], 8)

add_heading(doc, "3.3 Use Case Catalog", 2)
use_cases = [
    ("UC-01", "Sign in", "All users", "Authenticate using an active account."),
    ("UC-02", "Change password", "All users", "Verify current password and save a compliant password."),
    ("UC-03", "View personal profile", "All users", "View identity, role and assigned port."),
    ("UC-04", "Configure local notifications", "All users", "Enable website popup/voice and select minimum popup severity."),
    ("UC-05", "View operational dashboard", "All users", "View weather, risk, mode, zones, alerts and trends within access scope."),
    ("UC-06", "View alert center", "All users", "View role-scoped alerts and filter by port, zone, severity and date."),
    ("UC-07", "Hear High/Critical alert", "Authorized recipient", "Play Vietnamese speech for an accessible alert."),
    ("UC-08", "Acknowledge alert", "Admin, Port Manager", "Record that an authorized manager received an alert."),
    ("UC-09", "View response tasks", "All users", "View tasks allowed by role and assignment."),
    ("UC-10", "Assign response task", "Admin, Port Manager", "Assign a NEW task to an active Operator at the same port."),
    ("UC-11", "Acknowledge task", "Assigned Operator", "Confirm receipt of a personally assigned task."),
    ("UC-12", "Start task", "Assigned Operator", "Change an acknowledged task to in progress."),
    ("UC-13", "Complete task", "Assigned Operator", "Save a completion result and close an in-progress task."),
    ("UC-14", "Manage users", "Admin", "Search/filter, create and edit accounts with role-port rules."),
    ("UC-15", "Manage ports and zones", "Admin", "Create/update port and zone information."),
    ("UC-16", "Configure risk thresholds", "Admin, Port Manager", "View and update thresholds/zone overrides."),
    ("UC-17", "Configure SOP", "Admin, Port Manager", "Create, update or disable response rules."),
    ("UC-18", "Run simulation", "Admin", "Create/select a dataset and replay it through risk processing."),
    ("UC-19", "View simulation result", "Admin", "Review dangerous zones, map points and generated tasks."),
    ("UC-20", "Create five-day operating forecast", "Admin, Port Manager", "Convert weather forecast into daily risk and operating guidance."),
    ("UC-21", "Analyze long-range trend", "Admin, Port Manager", "Use PCA/K-Means risk analysis for 7/14/30/60/90-day planning."),
    ("UC-22", "Evaluate forecast accuracy", "Admin, Port Manager", "Compare forecast with actual observations when available."),
    ("UC-23", "Preview operational report", "Admin, Port Manager", "Query alerts, tasks or operation events by filter."),
    ("UC-24", "Export Excel/PDF", "Admin, Port Manager", "Download the previewed dataset and log the export."),
    ("UC-25", "View operation history", "All users", "Review role-scoped real and simulation events with friendly labels."),
]
add_table(doc, ["ID", "Use case", "Actors", "Summary"], use_cases, [1.7, 4.3, 3.2, 8], 8)

add_heading(doc, "4. Functional Requirements", 1)
add_heading(doc, "4.1 Authentication and Account Management", 2)
add_requirement(doc, "FR-AUTH-01", "Sign In", "User opens /login.", "Authenticate email/password and return the current user, role and assigned port in the session token.", "Valid active account → dashboard and role-based menu.", "Invalid credentials → generic error; inactive/deleted account → access denied; expired session → return to login.", "BR-01, BR-02, BR-03", "/login; POST /api/auth/login")
add_requirement(doc, "FR-AUTH-02", "Change Password", "Authenticated user opens account menu.", "Verify the current password; require confirmation and password policy compliance.", "Password is BCrypt-hashed and user is returned to sign in after success.", "Wrong current password or weak/mismatched new password → field-level message.", "BR-01, BR-02", "/change-password; PUT /api/auth/change-password")
add_requirement(doc, "FR-USR-01", "Manage Users", "Admin opens Người dùng.", "List, search and filter accounts by name/email/port, role and status; create/edit accounts; prevent deletion of an Admin account from the UI.", "Admin assigns one port when creating a Port Manager or Operator. Admin has no required port.", "Duplicate email, missing role/port or unauthorized caller → reject without changing data.", "BR-04, BR-05, BR-06", "/users, /users/new, /users/:id/edit; /api/users")

add_heading(doc, "4.2 Dashboard, Weather and Risk", 2)
add_requirement(doc, "FR-DASH-01", "Operational Overview", "User enters the protected application.", "Show current weather, final risk, operating mode, zone summary, active alerts and recent operational information.", "Data refreshes without requiring a full page reload and uses Vietnamese labels.", "Missing/stale data → explanatory state instead of technical null/code values.", "BR-07, BR-08", "/dashboard; GET /api/dashboard/summary")
add_requirement(doc, "FR-WEA-01", "Weather Ingestion", "Prefect schedule or authorized refresh executes.", "Collect OpenWeather data for configured active ports, normalize units and persist raw/normalized payload with timestamps.", "Successful reading becomes available to dashboard, risk evaluation and forecast evaluation.", "Provider/network failure → weather-fetch job is marked failed and previous valid data remains available.", "BR-07, BR-09", "Prefect weather collector; /api/weather/current|forecast|refresh")
add_requirement(doc, "FR-RISK-01", "Risk Evaluation", "A new reading or simulation snapshot is available.", "Evaluate wind, rainfall and visibility independently and use the highest factor as final risk.", "Store assessment, dominant factor, previous/final level and level-changed indicator.", "Invalid/missing optional factor → evaluate with available factors and explain missing data.", "BR-07, BR-08, BR-10", "Risk service/repository; /api/risk")
add_requirement(doc, "FR-RISK-02", "Threshold and SOP Configuration", "Admin or Port Manager opens configuration.", "View/update risk thresholds and zone overrides; manage response rules and their target operating mode/action.", "Validated changes apply to subsequent assessments and remain auditable.", "Invalid ranges, overlapping values or unauthorized role → reject.", "BR-10, BR-11", "/risk-config, /sop-rules; /api/risk, /api/sop-rules")

add_heading(doc, "4.3 Alert and Notification Management", 2)
add_requirement(doc, "FR-ALT-01", "Create and Scope Alerts", "Risk reaches HIGH or CRITICAL during operational or simulation processing.", "Create a port/zone alert containing weather context, severity and response recommendation. Scope access on the server.", "Admin sees all; Port Manager sees own port; Operator sees only alerts related to a personally assigned task.", "A user requesting an inaccessible alert receives not found/denied without data leakage.", "BR-12, BR-13, BR-14", "/alerts, /alerts/:id; GET /api/alerts")
add_requirement(doc, "FR-ALT-02", "Popup and Vietnamese Voice", "A new accessible HIGH/CRITICAL alert is detected by the client.", "Play an alert sound and Vietnamese speech containing port, zone, wind, rain, visibility, severity and response guidance.", "Popup follows the user's local severity preference and offers replay/acknowledge actions.", "Unsupported/blocked speech → preserve visual popup and replay control; lower severity remains in alert center.", "BR-12, BR-15", "Global alert panel; GET /api/alerts/:id/speech")
add_requirement(doc, "FR-ALT-03", "Email Port Manager", "A simulation creates a HIGH or CRITICAL alert.", "Send UTF-8 email to active Port Manager accounts assigned to that alert's port.", "Email identifies port, zone, severity, content, time and asks the manager to assign an Operator.", "SMTP failure is logged but must not roll back the completed simulation or alert record.", "BR-12, BR-16", "SMTP alert notifier")
add_requirement(doc, "FR-ALT-04", "Notification Preferences", "User opens notification settings.", "Enable/disable in-app popup and voice, and choose HIGH+CRITICAL or CRITICAL-only popup threshold.", "Preference is saved under the current user key on the current browser/device.", "Clearing browser storage or changing device resets local preferences; server email routing is unaffected.", "BR-15", "/notification-settings")

add_heading(doc, "4.4 Task Assignment and Execution", 2)
add_requirement(doc, "FR-TSK-01", "Generate Response Task", "A HIGH/CRITICAL scenario triggers a response.", "Create a NEW task with port, optional zone, priority, due time, title and response description.", "Task is visible to Admin and responsible Port Manager before assignment.", "Duplicate scenario steps shall be associated with their session/alert context for traceability.", "BR-17", "/tasks; operational.tasks")
add_requirement(doc, "FR-TSK-02", "Assign Task", "Admin or responsible Port Manager opens an unassigned NEW task.", "List only active Operators at the applicable port and save the selected assignee/due time.", "Selected Operator receives UTF-8 assignment email and then gains access to the related task and alert.", "Cross-port/inactive/non-Operator assignment or non-NEW task → reject. Email failure does not undo assignment.", "BR-13, BR-17, BR-18", "/alerts/:id or /tasks/:id; PATCH /api/tasks/:id/assignment")
add_requirement(doc, "FR-TSK-03", "Operator Workflow", "Assigned Operator opens the task.", "Enforce NEW → ACKNOWLEDGED → IN_PROGRESS → COMPLETED and require a completion note of at least 10 characters.", "Every state transition and result note is stored and written to operation history.", "Unassigned/wrong Operator, skipped state or invalid completion note → reject.", "BR-18, BR-19", "/tasks/:id; acknowledge/start/complete endpoints")

add_heading(doc, "4.5 Simulation and Historical Replay", 2)
add_requirement(doc, "FR-SIM-01", "Manage Simulation Dataset", "Admin opens Mô phỏng.", "Create, update, select or delete a dataset containing ordered weather snapshots for a port/zone.", "Validate snapshot order and measurements before saving.", "Empty/invalid dataset or unknown port/zone → show a clear validation error.", "BR-20", "/simulation; /api/simulation/datasets")
add_requirement(doc, "FR-SIM-02", "Replay Dataset", "Admin starts a selected dataset.", "Replay snapshots through weather, risk, operating mode, alert and task generation; keep simulation records separated by session ID.", "Show progress, friendly event feed, map/dangerous zones and result summary.", "Failure preserves completed transaction boundaries and reports the failed run.", "BR-20, BR-21", "/simulation, /simulation-results; POST /api/simulation/run|run-demo")

add_heading(doc, "4.6 Operational Forecast and AI Planning", 2)
add_requirement(doc, "FR-FOR-01", "Five-Day Operational Forecast", "Admin/Port Manager selects a port.", "Transform OpenWeather five-day forecast into daily weather risk, recommended operating mode and human-readable actions.", "Show aligned daily cards, safe activities, restricted activities and explanatory summary.", "No provider data → show error/empty state and do not invent a real observation.", "BR-22", "/forecast-planning; /api/weather/forecast, /api/simulation/forecast-plan")
add_requirement(doc, "FR-AI-01", "Long-Range Trend Analysis", "Admin/Port Manager selects 7, 14, 30, 60 or 90 days.", "Extend the available forecast into a planning trend and analyze risk using PCA scoring and K-Means cluster classification.", "Display risk trend, confidence that decreases with horizon, cluster explanation and operational recommendation.", "Output beyond OpenWeather coverage must be clearly described as planning support, not an official weather forecast.", "BR-23", "/ai-long-range-forecast; POST /api/ml/forecast-risk-analysis")
add_requirement(doc, "FR-EVAL-01", "Forecast Evaluation", "Admin/Port Manager opens evaluation.", "Compare planned wind/rain/visibility/risk with actual readings after the forecast time and calculate errors/status.", "Rows without actual observations remain Waiting for actual data; available rows are reconciled.", "Machine downtime/missing historical observation → report missing data explicitly; optional backfill may provide demo data but must be labeled.", "BR-24", "/forecast-evaluation; /api/forecast-evaluation")

add_heading(doc, "4.7 Reporting and Audit", 2)
add_requirement(doc, "FR-REP-01", "Preview Report", "Admin/Port Manager opens Báo cáo vận hành.", "Choose Alerts, Tasks or Operation Events and filter by port, zone, risk level and date range.", "Show total count and representative rows before export. Port Manager scope is fixed to own port.", "Invalid date/range or no data → explanatory message and disabled export.", "BR-25, BR-26", "/reports; GET /api/reports/preview")
add_requirement(doc, "FR-REP-02", "Export Excel/PDF", "User has a non-empty valid preview.", "Generate an XLSX table or formatted PDF with title, scope, filters, creation time and selected rows.", "Download filename is deterministic and export action is recorded in operation history.", "Unauthorized port scope or empty result → no blank file is generated.", "BR-25, BR-26", "/reports; GET /api/reports/export/xlsx|pdf")
add_requirement(doc, "FR-LOG-01", "Operation History", "User opens Lịch sử vận hành.", "Display role-scoped actual or simulation events with Vietnamese friendly event names and descriptions.", "Allow switching real/simulation scope and inspect events belonging to the same run.", "Unknown technical event type → safe friendly fallback rather than raw database code.", "BR-27", "/operation-log; GET /api/operation-events")

add_heading(doc, "5. Data Requirements", 1)
add_heading(doc, "5.1 Operational Entities", 2)
entities = [
    ("ports", "Port identity, coordinates, timezone, current mode/risk and active status."),
    ("users", "Email, BCrypt password hash, role, status and assigned_port_id."),
    ("refresh_tokens", "Hashed refresh/session token metadata."),
    ("zones", "Operational areas belonging to a port."),
    ("zone_threshold_overrides", "Optional zone-specific risk boundaries."),
    ("weather_readings", "Current, historical and simulation weather observations."),
    ("weather_fetch_jobs", "Outcome and timing of each weather collection attempt."),
    ("risk_thresholds", "Configurable boundaries by weather factor and risk level."),
    ("risk_assessments", "Per-reading final and factor risk evaluation."),
    ("sop_rules / sop_executions", "Response definition and execution history."),
    ("operation_mode_logs", "History of NORMAL/LIMITED/STOP transitions."),
    ("tasks", "Assignment, state lifecycle, due time and completion result."),
    ("alerts / alert_receipts", "Alert context and per-user read/acknowledgement state."),
    ("operation_events", "Append-oriented audit and simulation event stream."),
    ("simulation_datasets / sessions / snapshots", "Reusable input and isolated replay runs."),
]
add_table(doc, ["Entity", "Current responsibility"], entities, [5, 12])

add_heading(doc, "5.2 Analytics Entities", 2)
add_body(doc, "The analytics schema contains dimension tables for time, port, zone, risk and SOP action, plus fact tables for hourly weather, risk assessments, SOP executions, alerts and operation events. ETL watermarks prevent uncontrolled duplicate loads and support incremental refresh.")

add_heading(doc, "5.3 Data Retention and Integrity", 2)
add_bullets(doc, [
    "User removal is soft deletion; historical task, alert and event references remain traceable.",
    "All port-scoped operational records carry port_id; zone-scoped records additionally carry zone_id when applicable.",
    "Simulation records carry simulation_session_id and/or is_simulation to prevent confusion with actual operations.",
    "Passwords and secrets must never be stored in reports, logs or source control.",
    "Timestamps are stored with timezone and displayed in the operational local timezone.",
])

add_heading(doc, "6. Business Rules", 1)
rules = [
    ("BR-01", "Password must be at least eight characters and include uppercase, lowercase, number and special character."),
    ("BR-02", "Passwords are stored as BCrypt hashes; plaintext passwords are never logged."),
    ("BR-03", "JWT/session identity contains user ID, role and optional port ID used for authorization."),
    ("BR-04", "ADMIN is system-wide and does not require an assigned port."),
    ("BR-05", "PORT_MANAGER and OPERATOR must be assigned to one active port."),
    ("BR-06", "An Admin account cannot be deleted through the current user-management workflow."),
    ("BR-07", "Weather measurements preserve source and observation time; stale/missing data is shown explicitly."),
    ("BR-08", "Final risk is the maximum of wind, rain and visibility risk factors."),
    ("BR-09", "OpenWeather/API failure cannot erase the last valid operational reading."),
    ("BR-10", "Risk levels are LOW, MEDIUM, HIGH and CRITICAL; displayed Vietnamese labels are Thấp, Cần lưu ý, Cao and Rất cao."),
    ("BR-11", "Threshold and SOP changes apply to subsequent evaluations and are auditable."),
    ("BR-12", "Automatic popup/voice is limited to HIGH and CRITICAL according to user preference."),
    ("BR-13", "Port Manager can access alerts/tasks only for the assigned port."),
    ("BR-14", "Operator can access an alert only when a related task is assigned personally."),
    ("BR-15", "Website/voice preference is currently browser-local per user and device."),
    ("BR-16", "High/Critical alert email targets active Port Manager accounts of the alert's port."),
    ("BR-17", "A task may be assigned only to an active OPERATOR at the task's port."),
    ("BR-18", "Only the assigned Operator can acknowledge, start or complete a task."),
    ("BR-19", "Task transition order is NEW → ACKNOWLEDGED → IN_PROGRESS → COMPLETED; completion note minimum is 10 characters."),
    ("BR-20", "Simulation data is isolated by session and must be identifiable as simulation."),
    ("BR-21", "Simulation alerts/tasks use the same access and assignment rules as operational alerts/tasks."),
    ("BR-22", "Five-day operational planning derives from OpenWeather forecast data and risk thresholds."),
    ("BR-23", "Long-range output is decision-support trend analysis; confidence decreases with horizon."),
    ("BR-24", "Forecast evaluation cannot mark a row reconciled until an actual reading is available."),
    ("BR-25", "Port Manager may preview/export only the assigned port; Admin may select all ports."),
    ("BR-26", "Excel/PDF export must match the active filters and must not create an unexplained empty file."),
    ("BR-27", "Every material assignment, state transition, completion and report export is recorded in operation history."),
]
add_table(doc, ["ID", "Rule"], rules, [2, 15])

add_heading(doc, "7. External Interface Requirements", 1)
add_heading(doc, "7.1 User Interface", 2)
add_bullets(doc, [
    "Vietnamese operational wording shall be preferred over raw enum or event codes.",
    "Risk is consistently color-coded: LOW green, MEDIUM blue/amber, HIGH orange and CRITICAL red.",
    "Minimum readable body text should remain suitable for projector demonstration; empty/loading/error states are explicit.",
    "Desktop layout targets 1366×768 and above, with responsive behavior for narrower screens.",
    "Tables and cards must not concatenate identifiers and names; examples: “VNLCH — Cảng Lạch Huyện”.",
])
add_heading(doc, "7.2 Software Interfaces", 2)
add_table(doc, ["Interface", "Protocol / format", "Requirement"], [
    ("Frontend ↔ API", "HTTPS/HTTP REST, JSON", "Bearer-authenticated protected endpoints; UTF-8 content."),
    ("API ↔ PostgreSQL", "Npgsql / SQL", "Parameterized queries and least-privilege database account."),
    ("Prefect ↔ OpenWeather", "HTTPS REST/JSON", "Timeout, provider error handling and source metadata."),
    ("API ↔ SMTP", "SMTP with TLS", "UTF-8 alert/task messages; secrets supplied by environment variables."),
    ("API ↔ speech", "HTTP audio/mpeg", "Vietnamese speech only for accessible HIGH/CRITICAL alerts."),
    ("Metabase ↔ analytics", "PostgreSQL read-only", "No write permission to operational data."),
], [3.5, 4.5, 9])
add_heading(doc, "7.3 Deployment Interfaces", 2)
add_table(doc, ["Service", "Local port", "Health expectation"], [
    ("Frontend", "5173", "Returns HTTP 200."),
    ("Backend API", "5000", "GET /health returns HTTP 200."),
    ("PostgreSQL", "55432", "pg_isready succeeds."),
    ("Prefect", "4200", "Server and worker remain running."),
    ("Metabase", "3000", "API health check succeeds."),
], [5, 4, 8])

add_heading(doc, "8. Non-Functional Requirements", 1)
nfrs = [
    ("NFR-SEC-01", "Authorization", "Server-side queries enforce role and port/task scope; frontend hiding alone is insufficient."),
    ("NFR-SEC-02", "Secrets", "Database, JWT, OpenWeather and SMTP secrets reside in .env/runtime configuration and are never printed in reports."),
    ("NFR-SEC-03", "Input", "All SQL inputs are parameterized; identifiers and filters are validated."),
    ("NFR-PER-01", "Response", "Normal read APIs should complete within 2 seconds on the demonstration machine under expected load."),
    ("NFR-PER-02", "Refresh", "Dashboard/alert polling must not duplicate popups or speech for the same alert receipt."),
    ("NFR-REL-01", "Email isolation", "SMTP failure does not roll back an alert, simulation or saved task assignment."),
    ("NFR-REL-02", "Data continuity", "Provider downtime retains the latest valid reading and marks missing actual data clearly."),
    ("NFR-USA-01", "Language", "User-facing text is Vietnamese and avoids technical enum/event names."),
    ("NFR-USA-02", "Accessibility", "Critical information is conveyed by text and icon in addition to color; controls have accessible labels."),
    ("NFR-MNT-01", "Maintainability", "Frontend services, controllers and repositories remain separated; all builds compile with warnings treated visibly."),
    ("NFR-MNT-02", "Testing", "Key role, alert, task, user, forecast and report workflows have automated tests where supported by the environment."),
    ("NFR-DEP-01", "Deployment", "The complete application starts with Docker Compose using the documented app profile and health checks."),
]
add_table(doc, ["ID", "Attribute", "Requirement"], nfrs, [2.3, 3.2, 11.5])

add_heading(doc, "9. Requirement Traceability", 1)
add_table(doc, ["Requirement group", "Frontend", "API / repository", "Verification"], [
    ("Authentication/users", "/login, /users", "AuthController, UserController/UserRepository", "Frontend tests + backend build"),
    ("Alerts/voice", "/alerts, alert popup", "AlertController/AlertRepository, speech service", "Alert page/detail/service tests"),
    ("Alert email", "N/A", "SimulationController, SmtpAlertEmailNotifier", "Build + SMTP integration test"),
    ("Tasks", "/tasks, /tasks/:id", "TaskController/TaskRepository", "Task page/detail and API workflow tests"),
    ("Simulation", "/simulation", "SimulationController/Repository", "Run dataset and inspect session/result"),
    ("Forecast/AI", "/forecast-planning, /ai-long-range-forecast", "Weather/Simulation/Ml controllers", "Page/service tests and sample port run"),
    ("Evaluation", "/forecast-evaluation", "ForecastEvaluationController/Repository", "Preview and export tests"),
    ("Reports", "/reports", "ReportController/ReportRepository", "Preview, XLSX and PDF download tests"),
], [3.2, 4.7, 5.7, 3.4], 7.5)

add_heading(doc, "10. Application Message Catalog", 1)
messages = [
    ("MSG-01", "Đăng nhập thành công."),
    ("MSG-02", "Email hoặc mật khẩu không đúng."),
    ("MSG-03", "Không thể tải dữ liệu. Vui lòng thử lại."),
    ("MSG-04", "Không tìm thấy dữ liệu phù hợp với bộ lọc hiện tại."),
    ("MSG-05", "Đã lưu thay đổi."),
    ("MSG-06", "Người được chọn phải là nhân viên vận hành đang hoạt động và thuộc cùng cảng."),
    ("MSG-07", "Nhiệm vụ chưa được phân công."),
    ("MSG-08", "Hãy tiếp nhận nhiệm vụ trước khi bắt đầu thực hiện."),
    ("MSG-09", "Kết quả xử lý phải có ít nhất 10 ký tự."),
    ("MSG-10", "Đã hoàn tất nhiệm vụ và lưu kết quả xử lý."),
    ("MSG-11", "Chưa có dữ liệu thực tế để đối chiếu."),
    ("MSG-12", "Không tìm thấy dữ liệu báo cáo; hệ thống không tạo file trống."),
    ("MSG-13", "Đã tạo báo cáo và ghi vào lịch sử vận hành."),
    ("MSG-14", "Trình duyệt chưa cho phép phát giọng nói; bạn vẫn có thể xem nội dung cảnh báo."),
]
add_table(doc, ["Code", "Vietnamese message"], messages, [2.5, 14.5])

add_heading(doc, "11. Known Baseline Notes and Future Refinements", 1)
add_bullets(doc, [
    "Notification preferences are currently stored per browser/device. A future release may persist them server-side and add per-channel email/phone rules.",
    "Long-range 14–90 day values are planning trends derived from available forecast seeds and ML analysis; they are not official meteorological forecasts.",
    "Forgot-password and automated failed-login lockout were described by the older Report 3 but are not claimed as completed in this baseline.",
    "The report documents the intended secured simulation scope (Admin). Any API endpoint missing equivalent server authorization should be treated as a security hardening item, not as a granted permission.",
    "Screenshots should be refreshed after UI stabilization; route names and behavioral requirements in this document are the authoritative July 2026 baseline.",
])

doc.core_properties.title = "PORMS Report 3 - Software Requirements Specification"
doc.core_properties.subject = "Updated SRS aligned with July 2026 implementation"
doc.core_properties.author = "PORMS Project Team"
doc.core_properties.comments = "Generated from the implemented PORMS source baseline; original report preserved."

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
DOWNLOAD.write_bytes(OUTPUT.read_bytes())
print(OUTPUT)
print(DOWNLOAD)
