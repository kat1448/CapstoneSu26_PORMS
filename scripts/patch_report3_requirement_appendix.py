from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "docs" / "Report3_SRS_Updated_Same_Structure_20260731.docx"
DOWNLOAD = Path(r"C:\Users\LENOVO\Downloads\Report3_SRS_Updated_Same_Structure_20260731.docx")
DOWNLOAD_FALLBACK = Path(r"C:\Users\LENOVO\Downloads\Report3_SRS_Updated_Same_Structure_Appendix_Fixed_20260731.docx")


def set_cell(cell, value, bold=False):
    cell.text = str(value)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.bold = bold


def add_row(table, values):
    cells = table.add_row().cells
    for index, value in enumerate(values):
        set_cell(cells[index], value)


def remove_row(table, row):
    table._tbl.remove(row._tr)


def find_table(document, headers):
    expected = [header.casefold() for header in headers]
    for table in document.tables:
        if not table.rows:
            continue
        actual = [cell.text.strip().casefold() for cell in table.rows[0].cells]
        if actual == expected:
            return table
    raise RuntimeError(f"Could not find table with headers: {headers}")


document = Document(REPORT)

# Remove rows that an earlier generator accidentally appended to screen-layout
# tables because it relied on fixed table indexes.
for table in document.tables:
    if len(table.columns) != 7:
        continue
    for row in list(table.rows[1:]):
        first = row.cells[0].text.strip()
        second = row.cells[1].text.strip()
        if first.startswith("BR-") or (second.startswith("MSG-") and first.isdigit()):
            remove_row(table, row)

business_rules = find_table(document, ["#", "Rule Definition"])
common_requirements = find_table(document, ["#", "Requirement", "Description"])
messages = find_table(document, ["#", "Code", "Type", "Context", "Content"])

# Make the operation idempotent so the script can be rerun safely.
for row in list(business_rules.rows[1:]):
    if row.cells[0].text.strip() in {f"BR-{number}" for number in range(23, 35)}:
        remove_row(business_rules, row)

for row in list(common_requirements.rows[1:]):
    if row.cells[0].text.strip() in {str(number) for number in range(9, 16)}:
        remove_row(common_requirements, row)

for row in list(messages.rows[1:]):
    if row.cells[1].text.strip() in {f"MSG-{number}" for number in range(31, 43)}:
        remove_row(messages, row)

for item in [
    ("BR-23", "The five-day operational forecast shall use available OpenWeather forecast data and the configured PORMS risk thresholds."),
    ("BR-24", "Long-range forecast output is planning support based on PCA and K-Means analysis. Confidence shall decrease as the selected horizon increases, and the output shall not be presented as an official meteorological forecast."),
    ("BR-25", "A forecast evaluation record shall be marked as reconciled only when an actual weather observation exists. Missing actual data shall be stated clearly and shall not be replaced with zero values."),
    ("BR-26", "ADMIN may preview and export reports across accessible ports. PORT MANAGER may preview and export data only for the assigned port."),
    ("BR-27", "Excel and PDF exports shall contain the same data scope and filters shown in the report preview. The system shall not create an unexplained empty report."),
    ("BR-28", "A response task may be assigned only to an ACTIVE OPERATOR belonging to the same port as the alert. Only the assigned OPERATOR may process that task."),
    ("BR-29", "The task workflow shall follow NEW -> ACKNOWLEDGED -> IN_PROGRESS -> COMPLETED. Completion requires a result note of at least 10 characters."),
    ("BR-30", "Alert access shall be enforced by the server: ADMIN sees all accessible alerts, PORT MANAGER sees alerts of the assigned port, and OPERATOR sees only alerts linked to personally assigned tasks."),
    ("BR-31", "Automatic popup and Vietnamese voice notifications shall be triggered only for HIGH and CRITICAL alerts within the authenticated user's access scope."),
    ("BR-32", "A HIGH or CRITICAL alert shall notify active PORT MANAGER accounts of the affected port. Task-assignment email shall be sent only to the selected OPERATOR."),
    ("BR-33", "An email or speech-delivery failure shall be logged and shall not roll back the alert, simulation session, or task assignment."),
    ("BR-34", "Each successful report export shall record the actor, report type, active filters, scope, format, and export time in the operation history."),
]:
    add_row(business_rules, item)

for item in [
    ("9", "Role and Port Scope", "All alert, task, forecast, evaluation and report APIs shall enforce role and assigned-port scope on the server; hiding an item in the user interface alone is not sufficient authorization."),
    ("10", "Alert and Notification", "HIGH and CRITICAL alerts shall provide a clear Vietnamese popup and optional Vietnamese voice playback. Visual information shall remain available when browser audio is blocked or unavailable."),
    ("11", "Task Workflow", "The system shall restrict assignment to active same-port Operators, enforce valid status transitions, record processing timestamps, and require a completion result before a task is completed."),
    ("12", "Forecast Transparency", "Forecast screens shall identify the data source, distinguish forecast data from actual observations, translate technical risk codes into understandable Vietnamese, and explain the confidence and limitations of AI-assisted output."),
    ("13", "Report Integrity", "Report preview and exported Excel/PDF files shall use the same filters, access scope and source records. Files shall include report title, scope, selected period and creation time."),
    ("14", "Auditability", "Alert acknowledgement, task assignment and status changes, simulation runs, notification failures and report exports shall be recorded in operation history with actor and time information."),
    ("15", "Localization and Accessibility", "Operational labels, risk levels, event names and recommended actions shall be displayed in clear Vietnamese. Important information shall not rely on color or audio alone."),
]:
    add_row(common_requirements, item)

for item in [
    ("31", "MSG-31", "Toast", "Task assigned", "Đã phân công nhiệm vụ cho nhân viên vận hành."),
    ("32", "MSG-32", "Validation", "Invalid assignee", "Người được chọn phải đang hoạt động và thuộc cùng cảng với nhiệm vụ."),
    ("33", "MSG-33", "Validation", "Completion note", "Kết quả xử lý phải có ít nhất 10 ký tự."),
    ("34", "MSG-34", "Empty state", "Missing actual observation", "Chưa có dữ liệu thời tiết thực tế để đối chiếu."),
    ("35", "MSG-35", "Empty state", "Empty report", "Không có dữ liệu phù hợp với bộ lọc. Hệ thống không tạo báo cáo trống."),
    ("36", "MSG-36", "Toast", "Report exported", "Đã tạo báo cáo và ghi nhận hoạt động xuất file vào lịch sử vận hành."),
    ("37", "MSG-37", "Toast", "Alert acknowledged", "Đã xác nhận tiếp nhận cảnh báo."),
    ("38", "MSG-38", "Error", "Access denied by port scope", "Bạn không có quyền xem hoặc xử lý dữ liệu của cảng này."),
    ("39", "MSG-39", "Toast", "Task completed", "Đã hoàn tất nhiệm vụ và lưu kết quả xử lý."),
    ("40", "MSG-40", "Information", "Voice unavailable", "Trình duyệt chưa thể phát giọng đọc. Bạn vẫn có thể xem nội dung cảnh báo trên màn hình hoặc bấm Nghe cảnh báo."),
    ("41", "MSG-41", "Toast", "Notification preference saved", "Đã lưu tùy chọn nhận thông báo trên thiết bị này."),
    ("42", "MSG-42", "Error", "Notification delivery failed", "Không thể gửi một kênh thông báo. Cảnh báo và nhiệm vụ vẫn được lưu trong hệ thống."),
]:
    add_row(messages, item)

document.core_properties.comments = (
    "Original report structure preserved. Requirement Appendix updated for "
    "role/port scope, alert notification, task workflow, forecasts and reporting."
)
document.save(REPORT)
try:
    copyfile(REPORT, DOWNLOAD)
    download_output = DOWNLOAD
except PermissionError:
    # Microsoft Word locks an opened .docx. Keep the requested file intact and
    # provide a clearly named updated copy instead of asking the user to close it.
    copyfile(REPORT, DOWNLOAD_FALLBACK)
    download_output = DOWNLOAD_FALLBACK

print(REPORT)
print(download_output)
