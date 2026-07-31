from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SOURCE = Path(r"C:\Users\LENOVO\Downloads\Report3_SRS.docx")
ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Report3_SRS_Updated_Same_Structure_20260731.docx"
DOWNLOAD = Path(r"C:\Users\LENOVO\Downloads\Report3_SRS_Updated_Same_Structure_20260731.docx")


def set_cell(cell, value, bold=False):
    cell.text = str(value)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(8.5)
            r.bold = bold


def add_row(table, values):
    cells = table.add_row().cells
    for i, value in enumerate(values):
        set_cell(cells[i], value)
    return cells


def replace_paragraph(doc, startswith, new_text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            p.text = new_text
            p.style = "normal"
            return True
    return False


def find_paragraph(doc, exact):
    return next(p for p in doc.paragraphs if p.text.strip() == exact)


def move_before(element, anchor):
    anchor._p.addprevious(element)


def insert_paragraph(doc, anchor, text, style="normal"):
    p = doc.add_paragraph(text, style=style)
    move_before(p._p, anchor)
    return p


def insert_table(doc, anchor, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    # The original report uses localized/custom Word style names. Reuse its
    # existing table style id instead of introducing a new visual style.
    existing_style = doc.tables[7]._tbl.tblPr.tblStyle
    if existing_style is not None:
        table._tbl.tblPr.append(existing_style.__copy__())
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True)
    for row in rows:
        add_row(table, row)
    anchor._p.addprevious(table._tbl)
    return table


def insert_function_section(doc, anchor, number, title, trigger, description, normal_cases, abnormal_cases, rules, components):
    insert_paragraph(doc, anchor, f"{number} - {title}", "Heading 3")
    insert_paragraph(doc, anchor, f"Function trigger: {trigger}")
    insert_paragraph(doc, anchor, f"Function description: {description}")
    insert_paragraph(doc, anchor, "Screen layout:")
    insert_table(doc, anchor,
        ["#", "Component", "Comp. Type", "Editable", "Mandatory", "Default Value", "Description"],
        components)
    insert_paragraph(doc, anchor, "Function detail:")
    insert_paragraph(doc, anchor, "  Functionality:")
    insert_paragraph(doc, anchor, "    In Normal Cases:")
    for item in normal_cases:
        insert_paragraph(doc, anchor, item)
    insert_paragraph(doc, anchor, "    In Abnormal Cases:")
    for item in abnormal_cases:
        insert_paragraph(doc, anchor, item)
    insert_paragraph(doc, anchor, f"  Business rule: {rules}")
    insert_paragraph(doc, anchor, "––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––––")


doc = Document(SOURCE)

# Keep the original cover and formatting; update only the baseline date.
for p in doc.paragraphs:
    if "Danang, May 2026" in p.text:
        p.text = p.text.replace("Danang, May 2026", "Danang, July 2026")

# Existing tables retain their position and design.
change_log = doc.tables[0]
add_row(change_log, ["31/07/2026", "M, A", "Đinh Hải Quân", "Update SRS to the implemented July 2026 baseline: role/port-scoped alerts, task assignment lifecycle, SMTP email, Vietnamese voice alerts, operational forecast, AI long-range forecast, forecast evaluation, notification settings and Excel/PDF reports."])

actors = doc.tables[1]
set_cell(actors.rows[1].cells[2], "The System Administrator has system-wide access. Responsibilities include managing user accounts, ports and zones; viewing all accessible alerts and tasks; configuring risk thresholds and SOP rules; running simulations; using forecasting, evaluation, reporting and audit functions. ADMIN is not bound to a specific port.")
set_cell(actors.rows[2].cells[2], "The Port Manager is assigned to one port. The manager monitors that port, receives HIGH/CRITICAL alerts, acknowledges alerts, assigns NEW response tasks only to active Operators at the same port, supervises task progress and exports reports for the assigned port. The role cannot manage user accounts.")
set_cell(actors.rows[3].cells[2], "The Operator is assigned to one port and works within a personal task scope. The Operator sees alerts only when a related task is assigned personally, receives an assignment email, and processes the task through ACKNOWLEDGED, IN_PROGRESS and COMPLETED with a mandatory completion note.")

use_cases = doc.tables[2]
for row in use_cases.rows[1:]:
    if row.cells[1].text.strip() == "View Alert Panel":
        set_cell(row.cells[3], "ADMIN views all alerts; PORT MANAGER views alerts from the assigned port; OPERATOR views only alerts related to personally assigned tasks. HIGH/CRITICAL alerts can trigger popup and Vietnamese speech.")
    if row.cells[1].text.strip() == "View Task Log":
        set_cell(row.cells[1], "View and Process Tasks")
        set_cell(row.cells[3], "Users view tasks within role scope. Admin/Port Manager assigns a NEW task to an eligible same-port Operator; the assigned Operator acknowledges, starts and completes it with a result note.")
    if row.cells[1].text.strip() == "Start Simulation":
        set_cell(row.cells[2], "ADMIN")
        set_cell(row.cells[3], "Admin creates/selects a simulation dataset and replays historical or prepared weather snapshots through the operational pipeline.")
for row in [
    ("UC-38", "Configure Notification Preference", "All Actors", "Users enable/disable in-app popup and Vietnamese voice and select HIGH+CRITICAL or CRITICAL-only popup threshold on the current device."),
    ("UC-39", "Assign Response Task", "ADMIN, PORT MANAGER", "Authorized manager assigns a NEW task only to an active Operator belonging to the task's port."),
    ("UC-40", "Process Assigned Task", "OPERATOR", "Assigned Operator acknowledges, starts and completes a task with a completion result."),
    ("UC-41", "View Operational Forecast", "ADMIN, PORT MANAGER", "View five-day weather-based risk and operating recommendations for an accessible port."),
    ("UC-42", "View AI Long-range Forecast", "ADMIN, PORT MANAGER", "Analyze 7/14/30/60/90-day planning trends with PCA/K-Means risk analysis."),
    ("UC-43", "Evaluate Forecast", "ADMIN, PORT MANAGER", "Compare forecast values with actual weather readings when actual data is available."),
    ("UC-44", "Preview and Export Report", "ADMIN, PORT MANAGER", "Filter alerts, tasks or operation events, preview the result, then export Excel or PDF within access scope."),
]:
    add_row(use_cases, row)

screens = doc.tables[3]
for row in [
    ("26", "Task Workflow", "Task Detail Screen", "Allows assignment by Admin/Port Manager and acknowledgement, start and completion by the personally assigned Operator."),
    ("27", "Notification", "Notification Settings Screen", "Configures website popup, Vietnamese voice and minimum automatic-notification severity for the current browser/device."),
    ("28", "Forecast", "Operational Forecast Screen", "Shows five-day weather risk, recommended operating mode and human-readable response actions."),
    ("29", "Forecast", "AI Long-range Forecast Screen", "Shows planning trends from 7 days to 3 months, risk score, cluster, confidence and recommendation."),
    ("30", "Forecast", "Forecast Evaluation Screen", "Compares forecast and actual wind/rain/visibility/risk and explains missing actual data."),
    ("31", "Reporting", "Operational Report Screen", "Previews and exports filtered Alerts, Tasks or Operation Events to Excel/PDF."),
]:
    add_row(screens, row)

authorization = doc.tables[4]
for row in authorization.rows[1:]:
    name = row.cells[0].text.strip()
    if name == "Task Log Screen":
        set_cell(row.cells[1], "X (all accessible)"); set_cell(row.cells[2], "X (own port)"); set_cell(row.cells[3], "X (assigned only)")
    if name in ("Simulation Control Screen", "Simulation Results Screen"):
        set_cell(row.cells[2], ""); set_cell(row.cells[3], "")
for row in [
    ("Task Detail Screen", "X", "X (own port)", "X (assigned only)"),
    ("Notification Settings Screen", "X", "X", "X"),
    ("Operational Forecast Screen", "X", "X (own port)", ""),
    ("AI Long-range Forecast Screen", "X", "X (own port)", ""),
    ("Forecast Evaluation Screen", "X", "X (own port)", ""),
    ("Operational Report Screen", "X (all ports)", "X (own port)", ""),
]:
    add_row(authorization, row)

non_screen = doc.tables[5]
for row in [
    ("9", "Alert System", "Port-scoped Alert Delivery", "HIGH/CRITICAL alerts are visible to Admin and the assigned Port Manager. An Operator gains access only after a related task is assigned personally."),
    ("10", "Notification", "Vietnamese Voice Alert", "Accessible HIGH/CRITICAL alerts provide audio/mpeg Vietnamese speech; visual popup remains available if browser audio is blocked."),
    ("11", "Notification", "Manager Alert Email", "After simulation creates a HIGH/CRITICAL alert, SMTP email is sent to active Port Manager accounts at that port. Mail failure does not roll back the simulation."),
    ("12", "Task", "Assignment Email", "After a task is assigned, SMTP email is sent only to the selected Operator."),
    ("13", "Forecast", "Risk Trend Analysis", "PCA score and K-Means cluster support long-range operational planning; output is not an official meteorological forecast."),
    ("14", "Reporting", "Excel/PDF Generation", "The API generates an export that matches preview filters and records the export in operation history."),
]:
    add_row(non_screen, row)

entities = doc.tables[6]
for row in entities.rows[1:]:
    entity = row.cells[1].text.strip()
    if entity == "users":
        set_cell(row.cells[2], "Stores user accounts: email, BCrypt password_hash, full_name, role (ADMIN|PORT_MANAGER|OPERATOR), status (ACTIVE|INACTIVE|LOCKED), assigned_port_id, last_login_at, deleted_at and timestamps.")
    elif entity == "task_logs":
        set_cell(row.cells[1], "tasks")
        set_cell(row.cells[2], "Stores task_code, alert/port/zone/session references, title, description, priority, status, assigned_user_id, acknowledgement/start/completion timestamps, completion_note, due_at and audit timestamps.")
    elif entity == "alerts":
        set_cell(row.cells[2], "Stores port/zone/risk/session references, alert_type, severity, title, message, context JSONB, expiry and timestamps. Per-user state is stored separately in alert_receipts.")
    elif entity == "alert_recipients":
        set_cell(row.cells[1], "alert_receipts")
        set_cell(row.cells[2], "Stores one alert-user receipt with delivered_at, read_at and acknowledged_at. Access scope is enforced by role, assigned port and related task assignment.")
    elif entity == "simulation_sessions":
        set_cell(row.cells[2], "Records a replay run linked to simulation_dataset_id and port_id, with starter, status, progress/current step, generated counts and start/completion timestamps.")
    elif entity == "simulation_results":
        set_cell(row.cells[1], "simulation_datasets / simulation_snapshots")
        set_cell(row.cells[2], "Stores reusable scenario metadata and ordered weather snapshots. Run results are resolved from session-linked risk, alert, task and event records.")

# Product and old module descriptions are corrected in place.
replace_paragraph(doc, "The backend is built with", "The current implementation uses ASP.NET Core 10, PostgreSQL 16, Prefect 2.x, Metabase 0.50, Docker Compose, and a React 18/TypeScript frontend. Operational and analytics schemas remain separated in PostgreSQL.")
replace_paragraph(doc, "Function description: Displays the full paginated list of alerts", "Function description: Displays alerts within the authenticated user's server-enforced scope. ADMIN sees all ports; PORT MANAGER sees only the assigned port; OPERATOR sees only alerts connected to tasks assigned personally. Filters include port, zone, severity and date. Accessible HIGH/CRITICAL alerts can trigger popup, sound and Vietnamese speech according to local notification preferences.")
replace_paragraph(doc, "Clicking \"Đánh dấu đã đọc\"", "Clicking acknowledge calls PATCH /api/alerts/{id}/acknowledge and stores a per-user alert_receipts record. An inaccessible alert is not returned to the caller.")
replace_paragraph(doc, "\"Đánh dấu tất cả đã đọc\"", "The alert detail screen shows weather context, recommended action and related tasks. Admin/Port Manager may assign a NEW task within their permitted port scope.")
replace_paragraph(doc, "Statistics strip at top shows", "Statistics cards show waiting, CRITICAL, HIGH and acknowledged counts for the current access scope.")
replace_paragraph(doc, "Function description: Displays a chronological list (newest first) of auto-generated task recommendations", "Function description: Displays role-scoped operational tasks and supports their full lifecycle. ADMIN monitors accessible tasks; PORT MANAGER monitors and assigns tasks at the assigned port; OPERATOR sees only personally assigned tasks and can acknowledge, start and complete them with a result note.")
replace_paragraph(doc, "All tasks are read-only; no lifecycle management.", "Task workflow is NEW → ACKNOWLEDGED → IN_PROGRESS → COMPLETED. Completion requires a result note of at least 10 characters.")
replace_paragraph(doc, "Export CSV downloads filtered task data", "After assignment, only the selected Operator receives the assignment email and access to the related task/alert; other Operators at the port do not receive that personal notification.")
replace_paragraph(doc, "Function trigger: Triggered when ADMIN or PORT MANAGER navigates to the Operation Log", "Function trigger: Triggered when any authenticated user navigates to Lịch sử vận hành. Returned events are filtered by the user's role and accessible port/task scope.")
replace_paragraph(doc, "Function description: Displays the append-only audit trail", "Function description: Displays the append-oriented operation event history using friendly Vietnamese labels. Users can switch between actual and simulation history; technical event codes are translated before display. No edit or delete action is provided.")
replace_paragraph(doc, "Function trigger: Triggered when ADMIN or PORT MANAGER navigates to the Simulation section", "Function trigger: Triggered when ADMIN navigates to the Simulation section from the sidebar.")
replace_paragraph(doc, "Function description: Allows ADMIN and PORT MANAGER to upload a JSON dataset", "Function description: Allows ADMIN to create/select a reusable simulation dataset and replay ordered historical or prepared weather snapshots through Weather → Risk → Mode → Alert → Task. Generated data is isolated by simulation_session_id and displayed as simulation history.")
replace_paragraph(doc, "Scenario Name: Required.", "Dataset name and port are required. The dataset must contain at least one valid ordered weather snapshot; each snapshot includes wind, rainfall, visibility and optional zone context.")
replace_paragraph(doc, "Weather Snapshots JSON: Required", "Dataset snapshots are validated before save/run. The current screen supports saved datasets and a prepared demonstration sequence rather than claiming an implemented drag-and-drop-only workflow.")
replace_paragraph(doc, "Speed Multiplier: Required", "The current implementation runs the stored sequence and reports session progress, generated alerts/tasks and final risk. Replay speed controls described in earlier drafts are not treated as completed unless present in the active UI.")
replace_paragraph(doc, "SI-07: Database communication", "SI-07: Database communication must use parameterized Npgsql queries to prevent SQL injection. The current infrastructure repositories do not rely on Entity Framework Core.")
replace_paragraph(doc, "The PORMS system must maintain ≥ 99% uptime", "The PORMS demonstration environment uses Docker Compose restart policies and health checks for PostgreSQL, backend and Metabase. Weather collection is orchestrated by Prefect; service failure must be visible and must not erase the latest valid operational data.")

# Insert new functional sections before the existing Non-Functional Requirements, preserving chapter order.
anchor = find_paragraph(doc, "4. Non-Functional Requirements")
insert_paragraph(doc, anchor, "3.12 Operational Forecast", "Heading 3")
insert_function_section(doc, anchor, "3.12.1", "Five-Day Operational Forecast",
    "Triggered when ADMIN or PORT MANAGER opens Dự báo vận hành and selects an accessible port.",
    "Uses OpenWeather five-day forecast and PORMS risk thresholds to show daily risk, recommended operating mode and plain-language actions.",
    ["Display five aligned daily cards with weather/risk summary.", "Translate technical risk and weather codes into Vietnamese.", "Separate activities that may continue from activities that should be limited."],
    ["If forecast data is unavailable, show a clear error/empty state and do not present generated values as actual observations."],
    "BR-23",
    [("1", "Port selector", "Dropdown", "Yes", "Yes", "Assigned/default port", "Selects an accessible port."), ("2", "Five-day cards", "Card grid", "No", "Yes", "N/A", "Daily weather, risk and operating guidance."), ("3", "Action groups", "Information panels", "No", "Yes", "N/A", "Plain-language recommended and restricted activities.")])

insert_paragraph(doc, anchor, "3.13 AI Long-range Forecast", "Heading 3")
insert_function_section(doc, anchor, "3.13.1", "Long-range Operational Trend",
    "Triggered when ADMIN or PORT MANAGER opens Dự báo dài hạn and chooses 7, 14, 30, 60 or 90 days.",
    "Builds a planning trend from available forecast data and analyzes risk using PCA scoring and K-Means cluster classification.",
    ["Show risk score, rule risk, cluster explanation, confidence and operating recommendation.", "Decrease displayed confidence as the planning horizon grows.", "Clearly state that long-range results support planning and do not replace official weather forecasts."],
    ["If source forecast or ML analysis fails, show a retryable error without retaining misleading stale output."],
    "BR-24",
    [("1", "Horizon selector", "Button group", "Yes", "Yes", "7 days", "Chooses planning horizon."), ("2", "Trend chart", "Line chart", "No", "Yes", "N/A", "Risk score and confidence over time."), ("3", "Analysis table", "Table", "No", "Yes", "N/A", "Friendly risk, cluster and recommendation labels.")])

insert_paragraph(doc, anchor, "3.14 Forecast Evaluation", "Heading 3")
insert_function_section(doc, anchor, "3.14.1", "Forecast-versus-Actual Evaluation",
    "Triggered when ADMIN or PORT MANAGER opens Đánh giá dự báo.",
    "Compares forecast wind, rainfall, visibility and risk with actual weather observations after the planned time.",
    ["Show forecast/actual pair and error for every available factor.", "Mark a row reconciled only when actual data exists.", "Allow export of the evaluation dataset."],
    ["When the machine or collector was offline and no observation exists, show Chưa có dữ liệu thực tế instead of zero or a false comparison."],
    "BR-25",
    [("1", "Evaluation filters", "Filter panel", "Yes", "No", "Current scope", "Limits port and date range."), ("2", "Comparison rows", "Table", "No", "Yes", "N/A", "Forecast, actual and error values."), ("3", "Export", "Button", "No", "No", "Disabled without data", "Exports evaluation data.")])

insert_paragraph(doc, anchor, "3.15 Operational Reports", "Heading 3")
insert_function_section(doc, anchor, "3.15.1", "Preview and Export Report",
    "Triggered when ADMIN or PORT MANAGER opens Báo cáo vận hành.",
    "Previews filtered Alerts, Tasks or Operation Events and exports the same result to Excel or PDF.",
    ["Admin can choose all or one port; Port Manager is restricted to the assigned port.", "Select report type, port, zone, severity and date range, then preview.", "Export only a non-empty preview and record the action in operation history."],
    ["Invalid scope, date range or empty result produces an explanation and no blank file."],
    "BR-26, BR-27",
    [("1", "Report type", "Tabs", "Yes", "Yes", "Alerts", "Alerts, Tasks or Operation Events."), ("2", "Filters", "Filter panel", "Yes", "Yes", "Last 7 days", "Port, zone, risk and date range."), ("3", "Preview", "Table", "No", "Conditional", "Empty", "Rows matching active filters."), ("4", "Excel/PDF", "Buttons", "No", "Conditional", "Disabled", "Enabled only when preview contains data.")])

insert_paragraph(doc, anchor, "3.16 Notification Settings", "Heading 3")
insert_function_section(doc, anchor, "3.16.1", "Personal Notification Preference",
    "Triggered when any authenticated user opens Cài đặt thông báo.",
    "Controls website popup, Vietnamese voice and minimum automatic-popup severity for the current account on the current browser/device.",
    ["Enable/disable website popup.", "Enable/disable Vietnamese voice when popup is enabled.", "Choose HIGH+CRITICAL or CRITICAL-only automatic notification."],
    ["Browser audio restrictions keep the visual popup available. Clearing browser storage resets this local preference."],
    "BR-20, BR-21",
    [("1", "Website popup", "Switch", "Yes", "Yes", "Enabled", "Controls in-app alert popup."), ("2", "Vietnamese voice", "Switch", "Yes", "No", "Enabled", "Controls speech playback."), ("3", "Minimum severity", "Radio cards", "Yes", "Yes", "HIGH", "HIGH+CRITICAL or CRITICAL-only."), ("4", "Save", "Button", "No", "Yes", "N/A", "Stores preference for this user/device.")])

# Append current business rules and messages to the original appendices.
business_rules = doc.tables[26]
for row in [
    ("BR-23", "Five-day operational forecast uses OpenWeather forecast data and PORMS risk thresholds."),
    ("BR-24", "Long-range output is planning support based on PCA/K-Means analysis; confidence decreases with horizon and it is not an official weather forecast."),
    ("BR-25", "Forecast evaluation is reconciled only when an actual observation exists; missing actual data must be explained."),
    ("BR-26", "Port Manager may preview/export only the assigned port; Admin may export across accessible ports."),
    ("BR-27", "Excel/PDF export must match the active preview filters and must not create an unexplained empty file."),
    ("BR-28", "Only an active same-port Operator may be assigned a task; only that Operator may acknowledge, start and complete it."),
    ("BR-29", "Task state order is NEW → ACKNOWLEDGED → IN_PROGRESS → COMPLETED, with a completion note of at least 10 characters."),
]:
    add_row(business_rules, row)

messages = doc.tables[28]
for row in [
    ("31", "MSG-31", "Toast", "Task assigned", "Đã phân công nhiệm vụ cho nhân viên vận hành."),
    ("32", "MSG-32", "Validation", "Invalid assignee", "Người được chọn phải đang hoạt động và thuộc cùng cảng."),
    ("33", "MSG-33", "Validation", "Completion note", "Kết quả xử lý phải có ít nhất 10 ký tự."),
    ("34", "MSG-34", "Empty state", "Missing actual", "Chưa có dữ liệu thực tế để đối chiếu."),
    ("35", "MSG-35", "Empty state", "Empty report", "Không có dữ liệu phù hợp; hệ thống không tạo báo cáo trống."),
    ("36", "MSG-36", "Toast", "Report exported", "Đã tạo báo cáo và ghi vào lịch sử vận hành."),
]:
    add_row(messages, row)

# Ask Word to update TOC/page references when opened.
settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

doc.core_properties.title = "PORMS Report 3 - Software Requirement Specification (Updated)"
doc.core_properties.comments = "Original report structure, diagrams, images and tables preserved; content reconciled with July 2026 code baseline."
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
DOWNLOAD.write_bytes(OUTPUT.read_bytes())
print(OUTPUT)
print(DOWNLOAD)
